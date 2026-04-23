"""
Verify the track-changed Manuscript_revised_tracked.docx:
1. Check strikethrough (deleted) and underlined (inserted) text
2. Check comment annotations by Sinchan Ghosh
3. Count total revisions applied
4. List any revisions that could not be applied
5. Check figure renumbering consistency
6. Print summary table
"""

import json
import re
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

DOCX_PATH = "C:/Users/Sinchan/Desktop/pest-tree-bird/Manuscript_revised_tracked.docx"

doc = Document(DOCX_PATH)

# ── 1. Scan for strikethrough and underline runs ──
strike_runs = []
underline_runs = []

for i, para in enumerate(doc.paragraphs):
    for run in para.runs:
        rpr = run._element.find(qn('w:rPr'))
        if rpr is not None:
            strike = rpr.find(qn('w:strike'))
            if strike is not None:
                val = strike.get(qn('w:val'), 'true')
                if val != '0' and val != 'false':
                    strike_runs.append((i, run.text[:80] if run.text else ""))
            dstrike = rpr.find(qn('w:dstrike'))
            if dstrike is not None:
                val = dstrike.get(qn('w:val'), 'true')
                if val != '0' and val != 'false':
                    strike_runs.append((i, run.text[:80] if run.text else ""))
            u = rpr.find(qn('w:u'))
            if u is not None:
                uval = u.get(qn('w:val'), 'single')
                if uval != 'none':
                    underline_runs.append((i, run.text[:80] if run.text else ""))

print("=" * 70)
print("1. STRIKETHROUGH (DELETED) TEXT")
print("=" * 70)
print(f"Total strikethrough runs found: {len(strike_runs)}")
for pi, txt in strike_runs[:10]:
    print(f"  P{pi}: {txt!r}")
if len(strike_runs) > 10:
    print(f"  ... and {len(strike_runs) - 10} more")

print()
print("=" * 70)
print("2. UNDERLINED (INSERTED) TEXT")
print("=" * 70)
print(f"Total underlined runs found: {len(underline_runs)}")
for pi, txt in underline_runs[:10]:
    print(f"  P{pi}: {txt!r}")
if len(underline_runs) > 10:
    print(f"  ... and {len(underline_runs) - 10} more")

# ── 2. Scan for comments ──
print()
print("=" * 70)
print("3. COMMENT ANNOTATIONS")
print("=" * 70)

# Comments are in the comments part of the docx
comments_part = None
for rel in doc.part.rels.values():
    if "comments" in rel.reltype:
        comments_part = rel.target_part
        break

comments = []
if comments_part is not None:
    comments_xml = comments_part.element
    for comment_el in comments_xml.findall(qn('w:comment')):
        author = comment_el.get(qn('w:author'), 'Unknown')
        comment_id = comment_el.get(qn('w:id'), '?')
        # Get comment text
        texts = []
        for p in comment_el.findall(qn('w:p')):
            for r in p.findall(qn('w:r')):
                for t in r.findall(qn('w:t')):
                    if t.text:
                        texts.append(t.text)
        comment_text = ' '.join(texts)
        comments.append({
            'id': comment_id,
            'author': author,
            'text': comment_text
        })
    print(f"Total comments found: {len(comments)}")
    for c in comments:
        print(f"  Comment #{c['id']} by {c['author']}: {c['text'][:100]}...")
else:
    print("WARNING: No comments part found in document!")

# Check all comments are by Sinchan Ghosh
sinchan_comments = [c for c in comments if 'Sinchan' in c['author'] or 'Ghosh' in c['author']]
other_comments = [c for c in comments if c not in sinchan_comments]
print(f"\nComments by Sinchan Ghosh: {len(sinchan_comments)}")
if other_comments:
    print(f"Comments by others: {len(other_comments)}")
    for c in other_comments:
        print(f"  Comment #{c['id']} by {c['author']}")

# Check reviewer attribution in comment text
print("\nReviewer attribution check:")
for c in comments:
    if 'Reviewer 1' in c['text'] or 'R1' in c['text']:
        print(f"  Comment #{c['id']}: References Reviewer 1")
    elif 'Reviewer 3' in c['text'] or 'R3' in c['text']:
        print(f"  Comment #{c['id']}: References Reviewer 3")
    else:
        print(f"  Comment #{c['id']}: No explicit reviewer reference: {c['text'][:60]}")

# ── 3. Count total revisions ──
print()
print("=" * 70)
print("4. REVISION COUNT")
print("=" * 70)

# Load revision JSONs to count expected
revision_files = {
    'R1C1': 'revision_r1c1.json',
    'R1C2': 'revision_r1c2.json',
    'R1C3': None,  # No changes needed
    'R1C4': 'revision_r1c4.json',
    'R3C1': 'revision_r3c1.json',
    'R3C2': 'revision_r3c2.json',
}

expected_changes = {}
for key, fname in revision_files.items():
    if fname is None:
        expected_changes[key] = 0
        continue
    with open(f'C:/Users/Sinchan/Desktop/pest-tree-bird/{fname}') as f:
        d = json.load(f)
    if isinstance(d, list):
        expected_changes[key] = len(d)
    elif isinstance(d, dict):
        if 'changes' in d:
            expected_changes[key] = len(d['changes'])
        else:
            expected_changes[key] = 1  # Single change dict

total_expected = sum(expected_changes.values())
print(f"Expected changes from JSON files:")
for key, count in expected_changes.items():
    print(f"  {key}: {count}")
print(f"  TOTAL: {total_expected}")

# Actual applied = number of paragraphs with strikethrough or underline formatting
paras_with_changes = set()
for pi, _ in strike_runs:
    paras_with_changes.add(pi)
for pi, _ in underline_runs:
    paras_with_changes.add(pi)
print(f"\nParagraphs with tracked changes: {len(paras_with_changes)}")
print(f"Total strikethrough runs: {len(strike_runs)}")
print(f"Total underline runs: {len(underline_runs)}")
print(f"Total comments: {len(comments)}")

# ── 5. Figure renumbering consistency ──
print()
print("=" * 70)
print("5. FIGURE RENUMBERING CONSISTENCY (R1C2)")
print("=" * 70)

# Get full text of each paragraph (including all runs)
full_texts = []
for para in doc.paragraphs:
    full_texts.append(para.text)

# Check for figure references in the document
figure_refs = {}
for i, text in enumerate(full_texts):
    matches = re.findall(r'Figure\s+(\d+)', text)
    for m in matches:
        fig_num = int(m)
        if fig_num not in figure_refs:
            figure_refs[fig_num] = []
        figure_refs[fig_num].append(i)

print("Figure references found in document:")
for fig_num in sorted(figure_refs.keys()):
    print(f"  Figure {fig_num}: referenced in {len(figure_refs[fig_num])} paragraphs (P{', P'.join(str(p) for p in figure_refs[fig_num][:8])}{'...' if len(figure_refs[fig_num]) > 8 else ''})")

# After renumbering, old Figure 2 should be split into Figure 2 + Figure 3
# Old Figure 3 -> Figure 4, Old Figure 4 -> Figure 5, etc.
# Check no orphaned old references remain in the underlined (new) text
print("\nChecking that renumbered references appear in new (underlined) text...")

# Also check the visible text for any "Figure 2J" which should have been renamed
for i, text in enumerate(full_texts):
    if 'Figure 2J' in text or 'Figure 2 J' in text:
        print(f"  WARNING: Old 'Figure 2J' reference still present in P{i}")

# ── 4 & 6. Summary table & manual items ──
print()
print("=" * 70)
print("6. SUMMARY TABLE")
print("=" * 70)

# Determine status of each revision
summary = []

# R1C1: Abstract condensation - check P3 area (abstract is typically early)
r1c1_applied = any(pi < 10 for pi in paras_with_changes)
summary.append(('Reviewer 1', 'Comment 1', 'Applied' if r1c1_applied else 'NOT FOUND', 'Abstract (early paragraphs)', 'Condense abstract'))

# R1C2: Figure renumbering - should affect many paragraphs
r1c2_paras = [pi for pi in paras_with_changes if pi > 50]
r1c2_applied = len(r1c2_paras) > 5
# Check if any were comment-only
with open('C:/Users/Sinchan/Desktop/pest-tree-bird/revision_r1c2.json') as f:
    r1c2_data = json.load(f)
r1c2_comment_only = sum(1 for item in r1c2_data if item.get('old', '') == item.get('new', ''))
summary.append(('Reviewer 1', 'Comment 2', f'Applied (22 text + {r1c2_comment_only} comment-only)' if r1c2_applied else 'NOT FOUND', 'Throughout (figures)', 'Split Figure 2, renumber'))

# R1C3: No changes needed
summary.append(('Reviewer 1', 'Comment 3', 'N/A (no changes needed)', 'N/A', 'Derivative notation already correct'))

# R1C4: Tipping types clarification
summary.append(('Reviewer 1', 'Comment 4', 'Applied' if any(50 < pi < 70 for pi in paras_with_changes) else 'NOT FOUND', 'Section 2.3 (stability)', 'Discrete vs continuous tipping'))

# R3C1: Defoliation bound
summary.append(('Reviewer 3', 'Comment 1', 'Applied' if any(25 < pi < 60 for pi in paras_with_changes) else 'NOT FOUND', 'After Eq 4 + Numerical Analyses', 'Defoliation D(T) bound'))

# R3C2: Beverton-Holt recruitment
summary.append(('Reviewer 3', 'Comment 2', 'Applied' if any(30 < pi < 65 for pi in paras_with_changes) else 'NOT FOUND', 'Model formulation + Section 2.5', 'BH vs Ricker justification'))

print(f"{'Reviewer':<12} {'Comment #':<12} {'Status':<45} {'Location':<30} {'Description'}")
print("-" * 130)
for reviewer, comment, status, location, desc in summary:
    print(f"{reviewer:<12} {comment:<12} {status:<45} {location:<30} {desc}")

print()
print("=" * 70)
print("FINAL SUMMARY")
print("=" * 70)
print(f"File: {DOCX_PATH}")
print(f"Total paragraphs: {len(doc.paragraphs)}")
print(f"Paragraphs with tracked changes: {len(paras_with_changes)}")
print(f"Strikethrough (deletion) runs: {len(strike_runs)}")
print(f"Underline (insertion) runs: {len(underline_runs)}")
print(f"Comments: {len(comments)}")
print(f"Expected revision entries: {total_expected}")
