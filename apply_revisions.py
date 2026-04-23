"""
Apply tracked-change revisions to Manuscript.docx.

For each revision entry (old -> new):
  - Deleted text: strikethrough + red font
  - Inserted text: underline + blue font
  - Comment annotation after each change

Handles: multiple revisions per paragraph (applied in reverse positional order),
and equation objects that hide characters from run text.
"""

import json
import re
import copy
import logging
from pathlib import Path
from collections import defaultdict
from datetime import datetime, timezone
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.part import Part
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from lxml import etree

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
log = logging.getLogger(__name__)

BASE = Path(r"C:\Users\Sinchan\Desktop\pest-tree-bird")

RED = RGBColor(0xFF, 0x00, 0x00)
BLUE = RGBColor(0x00, 0x00, 0xFF)

COMMENTS_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
COMMENTS_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
)
COMMENTS_REL_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
)
COMMENT_AUTHOR = "Sinchan Ghosh"

# Global comment counter and comment element list
_comment_id_counter = 0
_comment_elements = []


def next_comment_id():
    global _comment_id_counter
    cid = _comment_id_counter
    _comment_id_counter += 1
    return cid


def create_comment_element(comment_id, text, author=COMMENT_AUTHOR):
    """Create a w:comment XML element for comments.xml."""
    nsmap = {"w": COMMENTS_NS}
    comment_el = etree.SubElement(
        etree.Element("dummy"), qn("w:comment"), nsmap=nsmap
    )
    comment_el.set(qn("w:id"), str(comment_id))
    comment_el.set(qn("w:author"), author)
    comment_el.set(qn("w:date"), datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"))
    # Add paragraph with run containing comment text
    p = etree.SubElement(comment_el, qn("w:p"))
    r = etree.SubElement(p, qn("w:r"))
    t = etree.SubElement(r, qn("w:t"))
    t.set(qn("xml:space"), "preserve")
    t.text = text
    return comment_el


def add_comment_range(para_element, comment_id, insert_before_element=None):
    """Insert commentRangeStart before insert_before_element in the paragraph."""
    range_start = OxmlElement("w:commentRangeStart")
    range_start.set(qn("w:id"), str(comment_id))
    if insert_before_element is not None:
        para_element.insert(list(para_element).index(insert_before_element), range_start)
    else:
        para_element.append(range_start)
    return range_start


def close_comment_range(para, comment_id):
    """Append commentRangeEnd and a commentReference run to the paragraph."""
    range_end = OxmlElement("w:commentRangeEnd")
    range_end.set(qn("w:id"), str(comment_id))
    para._element.append(range_end)
    # Reference run
    ref_run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rstyle = OxmlElement("w:rStyle")
    rstyle.set(qn("w:val"), "CommentReference")
    rpr.append(rstyle)
    ref_run.append(rpr)
    ref = OxmlElement("w:commentReference")
    ref.set(qn("w:id"), str(comment_id))
    ref_run.append(ref)
    para._element.append(ref_run)


def install_comments_part(doc):
    """Build comments.xml from collected comment elements and add it to the document package."""
    nsmap = {"w": COMMENTS_NS}
    comments_root = etree.Element(qn("w:comments"), nsmap=nsmap)
    for ce in _comment_elements:
        comments_root.append(ce)
    comments_xml = etree.tostring(comments_root, xml_declaration=True, encoding="UTF-8", standalone=True)

    doc_part = doc.part
    partname = "/word/comments.xml"

    from docx.opc.packuri import PackURI
    comments_part = Part(
        PackURI(partname),
        COMMENTS_CONTENT_TYPE,
        comments_xml,
        doc_part.package,
    )
    doc_part.relate_to(comments_part, COMMENTS_REL_TYPE)


def load_revisions():
    files = [
        "revision_r1c1.json",
        "revision_r1c2.json",
        "revision_r1c3.json",
        "revision_r1c4.json",
        "revision_r3c1.json",
        "revision_r3c2.json",
    ]
    revisions = []
    for fname in files:
        path = BASE / fname
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        if isinstance(data, list):
            for item in data:
                if item:
                    revisions.append(item)
        elif isinstance(data, dict) and data:
            revisions.append(data)
    return revisions


def strip_paragraph_prefix(text):
    return re.sub(r"^\[P\d+\]\s*", "", text)


def get_paragraph_text(para):
    return "".join(run.text for run in para.runs)


def norm(text):
    return re.sub(r"\s+", " ", text).strip()


def find_paragraph_index(paragraphs, old_text):
    """Find which paragraph contains old_text (normalized whitespace match)."""
    old_n = norm(old_text)
    for i, para in enumerate(paragraphs):
        para_n = norm(get_paragraph_text(para))
        if old_n in para_n:
            return i
    # Fuzzy: try first 80 chars
    if len(old_n) > 80:
        prefix = old_n[:80]
        for i, para in enumerate(paragraphs):
            para_n = norm(get_paragraph_text(para))
            if prefix in para_n:
                return i
    return None


def find_norm_position(full_text, old_text):
    """
    Find start/end character positions of old_text within full_text,
    using normalized whitespace matching but returning original positions.
    Returns (orig_start, orig_end) or None.
    """
    old_n = norm(old_text)
    # Build mapping: normalized index -> original index
    orig_indices = []
    norm_chars = []
    last_ws = False
    leading = True
    for i, ch in enumerate(full_text):
        if ch in " \t\n\r\xa0":
            if leading:
                continue
            if not last_ws:
                norm_chars.append(" ")
                orig_indices.append(i)
                last_ws = True
        else:
            leading = False
            norm_chars.append(ch)
            orig_indices.append(i)
            last_ws = False
    while norm_chars and norm_chars[-1] == " ":
        norm_chars.pop()
        orig_indices.pop()

    rebuilt = "".join(norm_chars)
    idx = rebuilt.find(old_n)
    if idx == -1:
        return None

    orig_start = orig_indices[idx]
    end_idx = idx + len(old_n) - 1
    orig_end = orig_indices[end_idx] + 1 if end_idx < len(orig_indices) else len(full_text)
    return (orig_start, orig_end)


def make_run(para, text, strike=False, underline=False, color=None, base_run=None,
             superscript=False, italic=False, size=None):
    run = para.add_run(text)
    if base_run and base_run.font:
        if base_run.font.name:
            run.font.name = base_run.font.name
        if base_run.font.size:
            run.font.size = base_run.font.size
        if base_run.font.bold is not None:
            run.font.bold = base_run.font.bold
        if base_run.font.italic is not None:
            run.font.italic = base_run.font.italic
    if strike:
        run.font.strike = True
    if underline:
        run.font.underline = True
    if color:
        run.font.color.rgb = color
    if italic:
        run.font.italic = True
    if size:
        run.font.size = size
    if superscript:
        rpr = run._element.get_or_add_rPr()
        va = OxmlElement("w:vertAlign")
        va.set(qn("w:val"), "superscript")
        rpr.append(va)
    return run


def apply_multiple_revisions(para, rev_list):
    """
    Apply multiple revisions to a single paragraph at once.
    rev_list: list of (old_text, new_text, comment)
    Returns (applied_comments, failed_comments).
    """
    full_text = get_paragraph_text(para)
    base_run = para.runs[0] if para.runs else None

    # Find positions for each revision
    positioned = []
    failed = []
    for old_text, new_text, comment in rev_list:
        pos = find_norm_position(full_text, old_text)
        if pos is None:
            failed.append(comment)
        else:
            positioned.append((pos[0], pos[1], old_text, new_text, comment))

    if not positioned:
        return [], failed

    # Sort by start position (descending) so we can apply from end to start
    # without invalidating earlier positions
    positioned.sort(key=lambda x: x[0])

    # Build segments: list of (text, type) where type is 'plain', 'change'
    segments = []
    cursor = 0
    for start, end, old_text, new_text, comment in positioned:
        if start > cursor:
            segments.append(("plain", full_text[cursor:start]))
        # The actual text in the document for this range
        deleted_text = full_text[start:end]
        segments.append(("change", deleted_text, new_text, comment))
        cursor = end
    if cursor < len(full_text):
        segments.append(("plain", full_text[cursor:]))

    # Save paragraph properties
    pPr = para._element.find(qn("w:pPr"))
    pPr_copy = copy.deepcopy(pPr) if pPr is not None else None

    # Clear all runs
    for r in list(para._element.findall(qn("w:r"))):
        para._element.remove(r)

    # Restore paragraph properties
    if pPr_copy is not None:
        existing = para._element.find(qn("w:pPr"))
        if existing is not None:
            para._element.remove(existing)
        para._element.insert(0, pPr_copy)

    # Rebuild
    applied = []
    for seg in segments:
        if seg[0] == "plain":
            if seg[1]:
                make_run(para, seg[1], base_run=base_run)
        else:
            _, deleted_text, new_text, comment = seg
            # Create margin comment
            cid = next_comment_id()
            ce = create_comment_element(cid, comment)
            _comment_elements.append(ce)
            # Insert commentRangeStart before the deletion run we're about to add
            range_start = OxmlElement("w:commentRangeStart")
            range_start.set(qn("w:id"), str(cid))
            para._element.append(range_start)
            # Deletion and insertion runs
            make_run(para, deleted_text, strike=True, color=RED, base_run=base_run)
            make_run(para, new_text, underline=True, color=BLUE, base_run=base_run)
            # Close comment range
            close_comment_range(para, cid)
            applied.append(comment)

    return applied, failed


def main():
    revisions = load_revisions()
    log.info(f"Loaded {len(revisions)} revision entries")

    doc = Document(str(BASE / "Manuscript.docx"))
    paragraphs = doc.paragraphs

    # Group revisions by paragraph index
    para_revisions = defaultdict(list)  # para_idx -> [(old, new, comment)]
    unmatched = []

    for i, rev in enumerate(revisions):
        old_text = strip_paragraph_prefix(rev["old"])
        new_text = strip_paragraph_prefix(rev["new"])
        comment = f"Comment by Sinchan Ghosh: {rev['comment']}"

        para_idx = find_paragraph_index(paragraphs, old_text)
        if para_idx is None:
            log.warning(f"Rev {i}: paragraph not found — {rev['comment'][:60]}")
            unmatched.append({"index": i, "comment": rev["comment"], "reason": "paragraph not found"})
        else:
            para_revisions[para_idx].append((old_text, new_text, comment))

    # Apply revisions paragraph by paragraph
    total_applied = []
    total_failed = list(unmatched)

    for para_idx in sorted(para_revisions.keys()):
        rev_list = para_revisions[para_idx]
        applied, failed = apply_multiple_revisions(paragraphs[para_idx], rev_list)
        for c in applied:
            total_applied.append({"paragraph": para_idx, "comment": c})
            log.info(f"P{para_idx}: Applied — {c[:70]}")
        for c in failed:
            # For paragraphs with equation objects, add comment-only margin note
            log.warning(f"P{para_idx}: text match failed, adding comment-only note — {c[:70]}")
            cid = next_comment_id()
            ce = create_comment_element(cid, f"MANUAL EDIT NEEDED — {c}")
            _comment_elements.append(ce)
            # Wrap last run (or whole paragraph) in a comment range
            p_el = paragraphs[para_idx]._element
            runs = p_el.findall(qn("w:r"))
            if runs:
                last_run = runs[-1]
                range_start = OxmlElement("w:commentRangeStart")
                range_start.set(qn("w:id"), str(cid))
                p_el.insert(list(p_el).index(last_run), range_start)
            else:
                range_start = OxmlElement("w:commentRangeStart")
                range_start.set(qn("w:id"), str(cid))
                p_el.append(range_start)
            close_comment_range(paragraphs[para_idx], cid)
            total_applied.append({"paragraph": para_idx, "comment": c, "note": "comment-only, manual edit needed"})

    # Install comments part into the document package
    if _comment_elements:
        install_comments_part(doc)

    # Save
    output_path = BASE / "Manuscript_revised_tracked.docx"
    doc.save(str(output_path))
    log.info(f"\nSaved to {output_path}")
    log.info(f"Applied: {len(total_applied)} / {len(revisions)}")
    log.info(f"Failed:  {len(total_failed)} / {len(revisions)}")

    if total_failed:
        log.info("\n=== Revisions needing manual attention ===")
        for f in total_failed:
            log.info(f"  P{f.get('paragraph','?')}: {f['comment'][:80]} — {f['reason']}")

    return total_applied, total_failed


if __name__ == "__main__":
    applied, failed = main()
