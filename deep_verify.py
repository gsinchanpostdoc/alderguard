"""
Deep verification of Manuscript_revised_tracked.docx:
1. Check comments.xml part exists with w:comment elements
2. Count comments and print author/text
3. For each paragraph with tracked changes, show partial (not whole-paragraph) strikethrough
4. Specifically check abstract paragraph for multiple small tracked changes
"""

import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.oxml.ns import qn

DOCX_PATH = "C:/Users/Sinchan/Desktop/pest-tree-bird/Manuscript_revised_tracked.docx"
doc = Document(DOCX_PATH)

# ── 1. Check comments.xml part exists ──
print("=" * 70)
print("1. COMMENTS.XML PART CHECK")
print("=" * 70)

comments_part = None
for rel in doc.part.rels.values():
    if "comments" in rel.reltype:
        comments_part = rel.target_part
        break

if comments_part is None:
    print("FAIL: No comments.xml part found!")
else:
    comments_xml = comments_part.element
    comment_elements = comments_xml.findall(qn('w:comment'))
    print(f"PASS: comments.xml exists with {len(comment_elements)} w:comment elements")

# ── 2. Count comments, print author/text ──
print()
print("=" * 70)
print("2. COMMENT DETAILS")
print("=" * 70)

for i, cel in enumerate(comment_elements):
    author = cel.get(qn('w:author'), 'Unknown')
    texts = []
    for p in cel.findall(qn('w:p')):
        for r in p.findall(qn('w:r')):
            for t in r.findall(qn('w:t')):
                if t.text:
                    texts.append(t.text)
    text = ' '.join(texts)
    print(f"  [{i}] Author: {author} | Text: {text[:120]}")

# ── 3 & 4. Check tracked-change paragraphs for partial vs whole-paragraph changes ──
print()
print("=" * 70)
print("3. PARTIAL vs WHOLE-PARAGRAPH TRACKED CHANGES")
print("=" * 70)

issues = []

for pi, para in enumerate(doc.paragraphs):
    runs = para.runs
    if not runs:
        continue

    has_strike = False
    has_underline = False
    has_normal = False
    strike_chars = 0
    underline_chars = 0
    normal_chars = 0

    for run in runs:
        rpr = run._element.find(qn('w:rPr'))
        is_strike = False
        is_underline = False
        if rpr is not None:
            s = rpr.find(qn('w:strike'))
            ds = rpr.find(qn('w:dstrike'))
            u = rpr.find(qn('w:u'))
            if s is not None and s.get(qn('w:val'), 'true') not in ('0', 'false'):
                is_strike = True
            if ds is not None and ds.get(qn('w:val'), 'true') not in ('0', 'false'):
                is_strike = True
            if u is not None and u.get(qn('w:val'), 'single') != 'none':
                is_underline = True

        txt_len = len(run.text) if run.text else 0
        if is_strike:
            has_strike = True
            strike_chars += txt_len
        elif is_underline:
            has_underline = True
            underline_chars += txt_len
        else:
            normal_chars += txt_len
            if txt_len > 0:
                has_normal = True

    if not (has_strike or has_underline):
        continue

    total_chars = strike_chars + underline_chars + normal_chars
    if total_chars == 0:
        continue

    # Check if this is a whole-paragraph replacement (no normal text at all)
    is_whole_para = not has_normal or normal_chars < 5

    status = "WHOLE-PARAGRAPH" if is_whole_para else "PARTIAL (good)"
    if is_whole_para and total_chars > 50:
        issues.append(pi)

    # Build a snippet showing the structure
    snippet_parts = []
    for run in runs:
        rpr = run._element.find(qn('w:rPr'))
        is_s = False
        is_u = False
        if rpr is not None:
            s = rpr.find(qn('w:strike'))
            ds = rpr.find(qn('w:dstrike'))
            u = rpr.find(qn('w:u'))
            if s is not None and s.get(qn('w:val'), 'true') not in ('0', 'false'):
                is_s = True
            if ds is not None and ds.get(qn('w:val'), 'true') not in ('0', 'false'):
                is_s = True
            if u is not None and u.get(qn('w:val'), 'single') != 'none':
                is_u = True

        txt = (run.text or '')[:40]
        if not txt.strip():
            continue
        if is_s:
            snippet_parts.append(f"~~{txt}~~")
        elif is_u:
            snippet_parts.append(f"__{txt}__")
        else:
            snippet_parts.append(txt)

    snippet = ' | '.join(snippet_parts[:8])
    if len(snippet_parts) > 8:
        snippet += ' | ...'

    print(f"  P{pi} [{status}] strike={strike_chars}ch, ins={underline_chars}ch, normal={normal_chars}ch")
    print(f"    Snippet: {snippet[:200]}")
    print()

# ── 5. Specific abstract check ──
print("=" * 70)
print("4. ABSTRACT PARAGRAPH (P7) DETAILED CHECK")
print("=" * 70)

abstract_para = doc.paragraphs[7]
run_types = []
for run in abstract_para.runs:
    rpr = run._element.find(qn('w:rPr'))
    is_s = False
    is_u = False
    if rpr is not None:
        s = rpr.find(qn('w:strike'))
        ds = rpr.find(qn('w:dstrike'))
        u = rpr.find(qn('w:u'))
        if s is not None and s.get(qn('w:val'), 'true') not in ('0', 'false'):
            is_s = True
        if ds is not None and ds.get(qn('w:val'), 'true') not in ('0', 'false'):
            is_s = True
        if u is not None and u.get(qn('w:val'), 'single') != 'none':
            is_u = True

    txt = run.text or ''
    if not txt.strip():
        continue

    if is_s:
        label = "DEL"
    elif is_u:
        label = "INS"
    else:
        label = "NORMAL"
    run_types.append((label, txt[:60]))

strike_count = sum(1 for t, _ in run_types if t == 'DEL')
ins_count = sum(1 for t, _ in run_types if t == 'INS')
normal_count = sum(1 for t, _ in run_types if t == 'NORMAL')

print(f"  Abstract has {len(run_types)} non-empty runs: {strike_count} DEL, {ins_count} INS, {normal_count} NORMAL")
if strike_count >= 3 and ins_count >= 3 and normal_count >= 3:
    print("  PASS: Multiple small tracked changes (not one giant deletion+insertion)")
elif strike_count <= 1 and ins_count <= 1:
    print("  FAIL: Only 1 deletion + 1 insertion = whole-paragraph replacement!")
else:
    print(f"  CHECK: {strike_count} deletions, {ins_count} insertions - may need review")

print()
print("  Run-by-run breakdown:")
for i, (label, txt) in enumerate(run_types):
    marker = {'DEL': '~~', 'INS': '__', 'NORMAL': '  '}[label]
    print(f"    [{i:2d}] {label:6s} {marker}{txt}{marker}")

# ── Summary ──
print()
print("=" * 70)
print("SUMMARY")
print("=" * 70)
if issues:
    print(f"ISSUES: {len(issues)} paragraph(s) appear to have whole-paragraph replacements:")
    for pi in issues:
        print(f"  - P{pi}")
    print("  These may need revision JSON fixes for more granular diffs.")
else:
    print("ALL GOOD: No whole-paragraph replacement issues detected.")
